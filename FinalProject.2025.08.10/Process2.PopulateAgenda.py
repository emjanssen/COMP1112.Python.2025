# - - - Imports - - - #

import os
from docx import Document

# - - - Template Inputs - - - #

#admin edits these variable values based on what they'd like their new edited agenda to say
newTitle        = "Council of Elrond"
newDate         = "2025/08/16"
newTime         = "13:30"
newLocation     = "A202"
newHost         = "Elrond"
newAttendees    = "Gimli, Boromir"
newAbsents      = "Celeborn, Thranduil"
newDelegates    = "Legolas Greenleaf"
newActionItems  = "1. Review Events To-Date\n2.Explore Options"
newPurpose      = "Discuss growing threat in Mordor."

# - - - Function Declarations - - - #

def createAgendaDirectory():

    agendaTemplateLocationDirectory = r"C:\Users\E.M. Janssen\Documents\GitHub\COMP1112.Python.2025\FinalProject.2025.08.10\Process2.AgendaTemplate"

    if not os.path.exists(agendaTemplateLocationDirectory):
        os.makedirs(agendaTemplateLocationDirectory)
        print("Directory created...")
    else:
        print("Directory already exists...")
    return agendaTemplateLocationDirectory

def openAgendaTemplate(agendaTemplateLocationDirectory):

    templateFileName = "AgendaTemplate.docx"
    templateFilePath = os.path.join(agendaTemplateLocationDirectory, templateFileName)

    if not os.path.exists(templateFilePath):
        print(f"File not found at: {templateFilePath}")
        return

    try:
        os.startfile(templateFilePath)
    except FileNotFoundError as exception:
        print(f"File not found: {templateFilePath}")
        return
    except Exception as exceptione:
        print(f"General error: {exception}")
        return

    try:
        agendaDocument = Document(templateFilePath)
    except Exception as exception:
        print(f"Error loading document... {exception}")
        return None

    return agendaDocument

# Validate Placeholder Functions #

def validateTitlePlaceholder(agendaDocument):

    searchWordTitle = "titlePlaceholder"

    for table in agendaDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                if searchWordTitle in cell.text:
                    print(f"Found '{searchWordTitle}' in table cell: {cell.text.strip()}")
                    return True, cell
    
    print(f"We did not find the placeholder: {searchWordTitle}.")
    return False, None

def validateDatePlaceholder(agendaDocument):

    searchWordDate = "datePlaceholder"

    for table in agendaDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                if searchWordDate in cell.text:
                    print(f"Found '{searchWordDate}' in table cell: {cell.text.strip()}.")
                    return True, cell
    
    print(f"We did not find the placeholder: {searchWordDate}.")
    return False, None

def validateTimePlaceholder(agendaDocument):

    searchWordTime = "timePlaceholder"

    for table in agendaDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                if searchWordTime in cell.text:
                    print(f"Found '{searchWordTime}' in table cell: {cell.text.strip()}.")
                    return True, cell
    
    print(f"We did not find the placeholder: {searchWordTime}.")
    return False, None

def validateLocationPlaceholder(agendaDocument):

    searchWordLocation = "locationPlaceholder"

    for table in agendaDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                if searchWordLocation in cell.text:
                    print(f"Found {searchWordLocation} in table cell: {cell.text.strip()}.")
                    return True, cell

    print(f"We did not find the placeholder: {searchWordLocation}.")
    return False, None

def validateHostPlaceholder(agendaDocument):

    searchWordHost = "hostPlaceholder"

    for table in agendaDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                if searchWordHost in cell.text:
                    print(f"Found '{searchWordHost}' in table cell: {cell.text.strip()}")
                    return True, cell

    print(f"We did not find the placeholder: {searchWordHost}.")
    return False, None

def validateAttendeesPlaceholder(agendaDocument):

    searchWordAttendees = "attendeesPlaceholder"

    for table in agendaDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                if searchWordAttendees in cell.text:
                    print(f"Found '{searchWordAttendees}' in table cell: {cell.text.strip()}")
                    return True, cell

    print(f"We did not find the placeholder: {searchWordAttendees}.")
    return False, None













# Replace Placeholder Functions #

def replaceTitlePlaceholder(agendaDocument, newTitle):

    titlePlaceholderFound, cell = validateTitlePlaceholder(agendaDocument)
    
    if titlePlaceholderFound:
        cell.text = newTitle
        print(f"Title replaced with: {newTitle}")
    else:
        print("Placeholder not replaced.")

    return agendaDocument
    
def replaceDatePlaceholder(agendaDocument, newDate):

    datePlaceholderFound, cell = validateDatePlaceholder(agendaDocument)
    
    if datePlaceholderFound:
        cell.text = newDate
        print(f"Date replaced with: {newDate}")
    else:
        print("Placeholder not replaced.")

    return agendaDocument

def replaceTimePlaceholder(agendaDocument, newTime):

    timePlaceholderFound, cell = validateTimePlaceholder(agendaDocument)
    
    if timePlaceholderFound:
        cell.text = newTime
        print(f"Time replaced with: {newTime}")
    else:
        print("Placeholder not replaced.")

    return agendaDocument

def replaceLocationPlaceholder(agendaDocument, newLocation):

    locationPlaceholderFound, cell = validateLocationPlaceholder(agendaDocument)

    if locationPlaceholderFound:
        cell.text = newLocation
        print(f"Location reaplced with: {newLocation}.")
    else:
        print("Placeholder not replaced.")
    
    return agendaDocument

def replaceHostPlaceholder(agendaDocument, newHost):

    hostPlaceholderFound, cell = validateHostPlaceholder(agendaDocument)
    
    if hostPlaceholderFound:
        cell.text = newHost
        print(f"Host replaced with: {newHost}")
    else:
        print("Placeholder not replaced.")

    return agendaDocument

def replaceAttendeesPlaceholder(agendaDocument, newAttendees):

    attendeesPlaceholderFound, cell = validateAttendeesPlaceholder(agendaDocument)
    
    if attendeesPlaceholderFound:
        cell.text = newAttendees
        print(f"Attendees replaced with: {newAttendees}")
    else:
        print("Placeholder not replaced.")

    return agendaDocument

# Save Modified Agenda #

def saveModifiedAgenda(agendaDocument):
    pathToModifiedAgenda = os.path.join(agendaTemplateLocationDirectory, "ModifiedAgenda.docx")
    agendaDocument.save(pathToModifiedAgenda)
    print(f"Modified agenda saved to: {pathToModifiedAgenda}")

# - - - Function Calls - - - #

agendaTemplateLocationDirectory = createAgendaDirectory()
agendaDocument = openAgendaTemplate(agendaTemplateLocationDirectory)

agendaDocument = replacePlaceholder(agendaDocument, "titlePlaceholder", newTitle)
agendaDocument = replacePlaceholder(agendaDocument, "datePlaceholder", newDate)
agendaDocument = replacePlaceholder(agendaDocument, "timePlaceholder", newTime)
agendaDocument = replacePlaceholder(agendaDocument, "locationPlaceholder", newLocation)
agendaDocument = replacePlaceholder(agendaDocument, "hostPlaceholder", newHost)
agendaDocument = replacePlaceholder(agendaDocument, "attendeesPlaceholder", newAttendees)
agendaDocument = replacePlaceholder(agendaDocument, "absentsPlaceholder", newAbsents)
agendaDocument = replacePlaceholder(agendaDocument, "delegatesPlaceholder", newDelegates)
agendaDocument = replacePlaceholder(agendaDocument, "actionItemsPlaceholder", newActionItems)
agendaDocument = replacePlaceholder(agendaDocument, "purposePlaceholder", newPurpose)

saveModifiedAgenda(agendaDocument)