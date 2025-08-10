# - - - Imports - - - #

import os
import docx
from docx import Document

# - - - Function Declarations - - - #

def createNotesCompilationDoc():

    meetingNotesDirectory = r"C:\Users\E.M. Janssen\Documents\GitHub\COMP1112.Python.2025\FinalProject.2025.08.10\Process3.MeetingNotes"

    if not os.path.exists(meetingNotesDirectory):
        os.makedirs(meetingNotesDirectory)
        print("Directory created...")
    else:
        print("Directory already exists...")
    
    return meetingNotesDirectory

def getTextFromMeetingNotes():

    meetingNotesOne = "MeetingNotes1.docx"
    meetingNotesTwo = "MeetingNotes2.docx"
    folderForUserFiles = r"C:\Users\E.M. Janssen\Documents\GitHub\COMP1112.Python.2025\FinalProject.2025.08.10\Process3.MeetingNotes"
    projectRootDirectory = r"C:\Users\E.M. Janssen\Documents\GitHub\COMP1112.Python.2025\FinalProject.2025.08.10"

    fileNameOne = os.path.join(projectRootDirectory, folderForUserFiles, meetingNotesOne)
    fileNameTwo = os.path.join(projectRootDirectory, folderForUserFiles, meetingNotesTwo)
    
    documentOne = docx.Document(fileNameOne)
    documentTwo = docx.Document(fileNameTwo)

    docOneText = []
    docTwoText = []

    for words in documentOne.paragraphs:
        docOneText.append(words.text)
    for words in documentTwo.paragraphs:
        docTwoText.append(words.text)

    return docOneText, docTwoText

def createCompilationDocument(docOneText, docTwoText, meetingNotesDirectory):
    
    compiledNotes = Document()
    compiledNotes.add_heading("Compiled Meeting Notes", 0)
    compiledNotes.add_heading("Team Member One")
    compiledNotes.add_paragraph(docOneText)
    compiledNotes.add_heading("Team Member Two")
    compiledNotes.add_paragraph(docTwoText)

    compiledNotesSavePath = os.path.join(meetingNotesDirectory, "CompiledMeetingNotes.docx")
    compiledNotes.save(compiledNotesSavePath)

# - - - Function Calls - - - #

meetingNotesDirectory = createNotesCompilationDoc()
docOneText, docTwoText = getTextFromMeetingNotes()
createCompilationDocument(docOneText, docTwoText, meetingNotesDirectory)